from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from docx.shared import Inches, Pt, RGBColor
from .models import Articles
from .forms import ArticlesFrom
from django.views.generic import DeleteView, CreateView, DetailView
from django.http import HttpResponse
from docx import Document
from django.contrib.auth import get_user_model
from django.contrib.auth.models import User
from django.db.models.signals import post_save
from django.dispatch import receiver
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side

@login_required
def applications_home(request):
    applications = Articles.objects.order_by('-date')
    return render(request, 'applications/applications_home.html', {'applications': applications})

class AplliDetailView(DetailView):
    model = Articles
    template_name = 'applications/appli_view.html'
    context_object_name = 'applications'

class AplliDeleteView(DeleteView):
    model = Articles
    success_url = '/applications/'
    template_name = 'applications/appli_delete.html'

class ArticlesCreateView(LoginRequiredMixin, CreateView):
    login_url = reverse_lazy('users:login')
    model = Articles
    template_name = 'applications/create.html'
    form_class = ArticlesFrom
    success_url = reverse_lazy('applications_home')

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.user = self.request.user
        self.object.save()
        return redirect('applications_home')

    # Заполнение значений из профиля пользователя
    @receiver(post_save, sender=Articles)
    def update_user(sender, instance, created, **kwargs):
        if created:
            instance.surname = instance.user.last_name
            instance.name = instance.user.first_name
            instance.mail = instance.user.email
            instance.middle_name = instance.user.middle_name
            instance.tab_number = instance.user.tab_number
            instance.phone_number = instance.user.phone_number
            instance.save()

# @login_required
# def create(request):
#     error = ''
#     if request.method == 'POST':
#         form = ArticlesFrom(request.POST)
#         if form.is_valid():
#             form.save()
#             return redirect('applications_home')
#         else:
#             error = 'Пожалуйста введите все значения!'
#
#     form = ArticlesFrom()
#     data = {
#         'form': form
#     }
#
#     return render(request, 'applications/create.html', data)

def export_to_docx(request, pk):
    article = get_object_or_404(Articles, pk=pk)

    doc = Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    doc_title = doc.add_paragraph('Данные по заявке', style='Heading 1')
    doc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc_title.paragraph_format.space_after = Inches(0.5)
    table = doc.add_table(rows=1, cols=13)
    table.autofit = True
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    header_texts = [
        'Департамент', 'Полное наименование отдела/сектора', 'Должность', 'Фамилия',
        'Имя', 'Отчество', 'Табельный номер', 'Адрес корпоративной почты',
        'Номер рабочего телефона', 'Услуга', 'Адрес установки оборудования',
        'Дата составления заявки', 'Комментарий'
    ]
    for i, text in enumerate(header_texts):
        cell = hdr_cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)

    data_row = table.add_row().cells
    data_values = [
        article.department, article.sector, article.post, article.surname,
        article.name, article.middle_name, article.tab_number, article.mail,
        article.phone_number, article.service, article.address,
        article.date.strftime('%d/%m/%Y'), article.komm,
    ]
    for i, value in enumerate(data_values):
        data_row[i].text = str(value)
        data_row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        data_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # doc.add_paragraph(f'Департамент: {article.department}')
    # doc.add_paragraph(f'Полное наименование отдела/сектора: {article.sector}')
    # doc.add_paragraph(f'Должность: {article.post}')
    # doc.add_paragraph(f'Фамилия: {article.surname}')
    # doc.add_paragraph(f'Имя: {article.name}')
    # doc.add_paragraph(f'Отчество: {article.middle_name}')
    # doc.add_paragraph(f'Табельный номер: {article.tab_number}')
    # doc.add_paragraph(f'Адрес корпоративной почты: {article.mail}')
    # doc.add_paragraph(f'Номер рабочего телефона: {article.phone_number}')
    # doc.add_paragraph(f'Услуга: {article.service}')
    # doc.add_paragraph(f'Адрес установки оборудования: {article.address}')
    # doc.add_paragraph(f'Дата составления заявки: {article.date}')
    # doc.add_paragraph(f'Комментарий: {article.komm}')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="article.docx"'
    doc.save(response)
    return response

def export_excel(request, pk):
    article = get_object_or_404(Articles, pk=pk)
    wb = Workbook()
    ws = wb.active

    ws.append(['Департамент', 'Полное наименование отдела/сектора', 'Должность', 'Фамилия', 'Имя', 'Отчество',
               'Табельный номер', 'Адрес корпоративной почты', 'Номер рабочего телефона', 'Услуга',
               'Адрес установки оборудования', 'Дата создания', 'Комментарий'])
    header_row = ws[1]
    for cell in header_row:
        cell.font = Font(bold=True)
    data_row = [article.department, article.sector, article.post, article.surname, article.name, article.middle_name,
               article.tab_number, article.mail, article.phone_number, article.service, article.address,
               article.date.strftime('%d/%m/%Y'), article.komm]
    ws.append(data_row)

    border_style = Side(style='thin', color='000000')
    border = Border(left=border_style, right=border_style, top=border_style, bottom=border_style)
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.border = border

    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2) * 1.2
        ws.column_dimensions[column].width = adjusted_width

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=article_{article.pk}.xlsx'
    wb.save(response)
    return response
